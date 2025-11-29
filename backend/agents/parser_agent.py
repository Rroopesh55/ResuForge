"""
Resume Parser Agent

This module provides the ParserAgent class which extracts text and layout metadata
from resume files in PDF and DOCX formats. The agent preserves structural information
including margins, fonts, and page dimensions to enable layout-aware resume optimization.

Author: ResuForge Team
Date: 2024
"""

import os
import logging
from typing import Dict, Any, List
import docx
from pypdf import PdfReader

try:
    from backend.utils.encoding_utils import safe_encode_text, detect_encoding, strip_control_characters
except ModuleNotFoundError:
    from utils.encoding_utils import safe_encode_text, detect_encoding, strip_control_characters

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ParserAgent:
    """
    Agent responsible for parsing resume files and extracting both content and layout metadata.
    
    Supports PDF and DOCX formats with layout-aware extraction to preserve formatting
    information necessary for the "Iron Rules" constraint system.
    """
    
    def __init__(self):
        """Initialize the Parser Agent."""
        logger.info("ParserAgent initialized")

    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parses a resume file (PDF or DOCX) and extracts text + layout metadata.
        
        Args:
            file_path (str): Absolute path to the resume file
            
        Returns:
            Dict[str, Any]: Dictionary containing:
                - text (str): Concatenated text content
                - metadata (dict): Layout information (margins, fonts, page count)
                - raw_content (list): List of paragraphs in original order
                
        Raises:
            ValueError: If file format is not supported
            
        Example:
            >>> parser = ParserAgent()
            >>> result = parser.parse_file("/path/to/resume.docx")
            >>> print(result['metadata']['page_count'])
            1
        """
        logger.info(f"Parsing file: {file_path}")
        detected_encoding = detect_encoding(file_path)
        logger.debug(f"Using detected encoding: {detected_encoding}")
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".docx":
            logger.debug("Detected DOCX format")
            return self._parse_docx(file_path)
        elif ext == ".pdf":
            logger.debug("Detected PDF format")
            return self._parse_pdf(file_path)
        else:
            logger.error(f"Unsupported file format: {ext}")
            raise ValueError(f"Unsupported file format: {ext}")

    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a DOCX file and extract text with layout metadata.
        
        This method uses python-docx to extract:
        - Paragraph text
        - Section margins
        - Font information
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            Dict[str, Any]: Parsed content with metadata
            
        Note:
            Page count is approximated using section count as DOCX doesn't
            directly expose page breaks without rendering.
        """
        logger.info(f"Parsing DOCX file: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            text_content = []
            layout_meta = {
                "page_count": len(doc.sections),  # Approximation
                "margins": [],
                "fonts": set()
            }

            # Extract Layout Metadata from sections
            logger.debug(f"Extracting layout from {len(doc.sections)} sections")
            for idx, section in enumerate(doc.sections):
                margin_data = {
                    "top": section.top_margin.inches if section.top_margin else None,
                    "bottom": section.bottom_margin.inches if section.bottom_margin else None,
                    "left": section.left_margin.inches if section.left_margin else None,
                    "right": section.right_margin.inches if section.right_margin else None,
                }
                layout_meta["margins"].append(margin_data)
                logger.debug(f"Section {idx} margins: {margin_data}")

            # Extract Text & Styles from paragraphs
            logger.debug(f"Extracting text from {len(doc.paragraphs)} paragraphs")
            for para in doc.paragraphs:
                if para.text.strip():
                    cleaned = self._clean_text(para.text)
                    text_content.append(cleaned)
                    # Attempt to extract font information
                    if para.style and hasattr(para.style, 'font'):
                        if para.style.font.name:
                            layout_meta["fonts"].add(para.style.font.name)

            # Convert set to list for JSON serialization
            layout_meta["fonts"] = list(layout_meta["fonts"])
            
            logger.info(f"Successfully parsed DOCX: {len(text_content)} paragraphs, "
                       f"{layout_meta['page_count']} sections")

            return {
                "text": "\n".join(text_content),
                "metadata": layout_meta,
                "raw_content": text_content  # List of paragraphs for easier reconstruction
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {str(e)}")
            raise

    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a PDF file and extract text content.
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            Dict[str, Any]: Parsed content with basic metadata
            
        Note:
            PDF parsing has limitations - exact margin extraction requires
            specialized PDF analysis. Current implementation focuses on text
            extraction and page count.
        """
        logger.info(f"Parsing PDF file: {file_path}")
        
        try:
            reader = PdfReader(file_path)
            text_content = []
            layout_meta = {
                "page_count": len(reader.pages),
                "margins": "Unknown (PDF)",  # Hard to extract without OCR/Vision
            }

            logger.debug(f"PDF has {layout_meta['page_count']} pages")
            
            for page_num, page in enumerate(reader.pages):
                page_text = self._clean_text(page.extract_text() or "")
                text_content.append(page_text)
                logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")

            logger.info(f"Successfully parsed PDF: {layout_meta['page_count']} pages")

            return {
                "text": "\n".join(text_content),
                "metadata": layout_meta,
                "raw_content": text_content
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF file: {str(e)}")
            raise

    def _clean_text(self, text: str) -> str:
        """
        Normalize text using the encoding utilities to support international resumes.
        """
        if not text:
            return ""
        encoded = safe_encode_text(text)
        stripped = strip_control_characters(encoded)
        return stripped.strip()

