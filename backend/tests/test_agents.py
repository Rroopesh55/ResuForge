"""
Tests for Backend Agents

This test suite validates the core functionality of all backend agents:
- ParserAgent: PDF/DOCX text extraction
- JDAgent: Job description keyword analysis
- RewriteAgent: AI-powered bullet rewriting
- LayoutEngine: DOCX editing and constraint validation

Run with: python -m pytest tests/test_agents.py -v
"""

import pytest
import os
from agents.parser_agent import ParserAgent
from agents.jd_agent import JDAgent
from agents.rewrite_agent import RewriteAgent
from agents.layout_engine import LayoutEngine


# ============================================================================
# ParserAgent Tests
# ============================================================================

class TestParserAgent:
    """Test resume text extraction from PDF and DOCX files"""
    
    def test_parser_initialization(self):
        """Test ParserAgent can be initialized"""
        parser = ParserAgent()
        assert parser is not None
    
    def test_extract_text_from_docx(self, sample_docx):
        """Test DOCX text extraction"""
        parser = ParserAgent()
        result = parser.parse_file(sample_docx)
        
        assert isinstance(result, dict)
        assert 'text' in result
        assert 'raw_content' in result
        assert len(result['text']) > 0
        # Check for common resume content
        text_lower = result['text'].lower()
        assert any(keyword in text_lower for keyword in ['experience', 'education', 'skills', 'engineer'])
    
    def test_extract_text_from_pdf(self, sample_pdf):
        """Test PDF text extraction"""
        parser = ParserAgent()
        result = parser.parse_file(sample_pdf)
        
        assert isinstance(result, dict)
        assert 'text' in result
        assert len(result['text']) > 0
    
    def test_invalid_file_path(self):
        """Test handling of nonexistent file"""
        parser = ParserAgent()
        
        with pytest.raises(Exception):
            parser.parse_file("nonexistent_file.docx")


# ============================================================================
# JDAgent Tests
# ============================================================================

class TestJDAgent:
    """Test job description analysis and keyword extraction"""
    
    def test_jd_agent_initialization(self):
        """Test JDAgent can be initialized"""
        jd_agent = JDAgent()
        assert jd_agent is not None
    
    def test_analyze_simple_jd(self):
        """Test keyword extraction from simple JD"""
        jd_agent = JDAgent()
        jd_text = """
        Looking for a Python developer with experience in FastAPI and Docker.
        Must have strong knowledge of REST APIs and database design.
        """
        
        result = jd_agent.analyze_jd(jd_text)
        
        assert 'keywords' in result
        assert isinstance(result['keywords'], list)
        assert len(result['keywords']) > 0
        
        # Check for obvious keywords
        keywords_lower = [k.lower() for k in result['keywords']]
        assert 'python' in keywords_lower
        assert 'fastapi' in keywords_lower or 'docker' in keywords_lower
    
    def test_analyze_empty_jd(self):
        """Test handling of empty JD"""
        jd_agent = JDAgent()
        result = jd_agent.analyze_jd("")
        
        # Should return structure even if empty
        assert 'keywords' in result
        assert isinstance(result['keywords'], list)
    
    def test_analyze_complex_jd(self):
        """Test extraction from realistic job description"""
        jd_agent = JDAgent()
        jd_text = """
        Senior Software Engineer - Backend
        
        We are seeking an experienced backend engineer to join our team.
        
        Requirements:
        - 5+ years of experience with Python
        - Strong knowledge of FastAPI, Django, or Flask
        - Experience with PostgreSQL and Redis
        - Familiarity with Docker and Kubernetes
        - AWS or GCP cloud experience
        
        Nice to have:
        - Machine learning experience
        - Microservices architecture
        """
        
        result = jd_agent.analyze_jd(jd_text)
        
        assert len(result['keywords']) >= 5  # Should extract multiple keywords


# ============================================================================
# LayoutEngine Tests
# ============================================================================

class TestLayoutEngine:
    """Test DOCX editing and constraint validation"""
    
    def test_layout_engine_initialization(self):
        """Test LayoutEngine can be initialized"""
        engine = LayoutEngine()
        assert engine is not None
    
    def test_validate_constraints_within_limit(self):
        """Test validation passes when within character limit"""
        engine = LayoutEngine()
        
        original = "Led team of 5 engineers"
        new_text = "Spearheaded cross-functional team of 5 engineers"
        max_chars = 60
        
        result = engine.validate_constraints(original, new_text, max_chars)
        assert result is True
    
    def test_validate_constraints_exceeds_limit(self):
        """Test validation fails when exceeding character limit"""
        engine = LayoutEngine()
        
        original = "Short text"
        new_text = "This is a very long piece of text that definitely exceeds the character limit"
        max_chars = 20
        
        result = engine.validate_constraints(original, new_text, max_chars)
        assert result is False
    
    def test_validate_constraints_exact_limit(self):
        """Test validation at exact character limit"""
        engine = LayoutEngine()
        
        original = "text"
        new_text = "12345"  # Exactly 5 chars
        max_chars = 5
        
        result = engine.validate_constraints(original, new_text, max_chars)
        assert result is True
    
    def test_reconstruct_docx(self, sample_docx, tmp_path):
        """Test DOCX reconstruction with text replacement"""
        engine = LayoutEngine()
        
        output_path = tmp_path / "output.docx"
        replacements = {
            "Engineer": "Senior Engineer"
        }
        
        engine.reconstruct_docx(
            original_path=sample_docx,
            output_path=str(output_path),
            replacements=replacements
        )
        
        # Verify output file was created
        assert output_path.exists()
        assert output_path.stat().st_size > 0


# ============================================================================
# RewriteAgent Tests
# ============================================================================

class TestRewriteAgent:
    """Test AI-powered bullet rewriting"""
    
    def test_rewrite_agent_initialization(self):
        """Test RewriteAgent can be initialized"""
        agent = RewriteAgent()
        assert agent is not None
    
    @pytest.mark.skip(reason="Requires Ollama running - integration test")
    def test_batch_rewrite_with_keywords(self):
        """Test bullet rewriting with keyword injection"""
        agent = RewriteAgent()
        
        bullets = ["Led team of 5 engineers"]
        keywords = ["cross-functional", "agile"]
        constraints = [{"max_chars": 100}]
        style = "bold"
        
        result = agent.batch_rewrite(bullets, keywords, constraints, style)
        
        assert isinstance(result, list)
        assert len(result) == 1
        assert len(result[0]) <= 100
        
        # Check if keywords were incorporated
        result_lower = result[0].lower()
        assert "cross-functional" in result_lower or "agile" in result_lower
    
    @pytest.mark.skip(reason="Requires Ollama running - integration test")
    def test_batch_rewrite_multiple_bullets(self):
        """Test rewriting multiple bullets"""
        agent = RewriteAgent()
        
        bullets = [
            "Built web applications",
            "Managed database systems"
        ]
        keywords = ["scalable", "cloud", "microservices"]
        constraints = [{"max_chars": 120}, {"max_chars": 120}]
        
        result = agent.batch_rewrite(bullets, keywords, constraints, "safe")
        
        assert len(result) == 2
        assert all(len(bullet) <= 120 for bullet in result)
    
    @pytest.mark.skip(reason="Requires Ollama running - integration test")
    def test_rewrite_styles(self):
        """Test different rewriting styles"""
        agent = RewriteAgent()
        
        bullets = ["Developed software applications"]
        keywords = ["Python", "FastAPI"]
        constraints = [{"max_chars": 100}]
        
        for style in ["safe", "bold", "creative"]:
            result = agent.batch_rewrite(bullets, keywords, constraints, style)
            assert len(result) == 1
            assert len(result[0]) > 0


# ============================================================================
# Fixtures
# ============================================================================

@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample DOCX file for testing"""
    from docx import Document
    
    doc = Document()
    doc.add_heading('Software Engineer', 0)
    doc.add_paragraph('Experience:')
    doc.add_paragraph('Led team of engineers at TechCorp')
    doc.add_paragraph('Built scalable web applications')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('Bachelor of Science in Computer Science')
    
    doc_path = tmp_path / "sample_resume.docx"
    doc.save(str(doc_path))
    
    return str(doc_path)


@pytest.fixture
def sample_pdf(tmp_path):
    """Create a sample PDF file for testing"""
    # For now, skip PDF creation or use a pre-existing test file
    # In a real scenario, you'd use reportlab or pypdf to create a test PDF
    pytest.skip("PDF fixture not yet implemented - use actual test file")


# ============================================================================
# Integration Tests
# ============================================================================

class TestEndToEndWorkflow:
    """Test complete workflow integration"""
    
    def test_parse_and_validate(self, sample_docx):
        """Test parsing a resume and validating edits"""
        parser = ParserAgent()
        engine = LayoutEngine()
        
        # Parse
        result = parser.parse_file(sample_docx)
        assert 'text' in result
        assert len(result['text']) > 0
        
        # Validate an edit
        original_bullet = "Led team of engineers"
        new_bullet = "Spearheaded cross-functional team of engineers"
        max_chars = len(original_bullet) + 20
        
        is_valid = engine.validate_constraints(original_bullet, new_bullet, max_chars)
        assert is_valid is True
    
    @pytest.mark.skip(reason="Requires Ollama - integration test")
    def test_full_optimization_pipeline(self, sample_docx):
        """Test complete optimization workflow"""
        parser = ParserAgent()
        jd_agent = JDAgent()
        rewrite_agent = RewriteAgent()
        engine = LayoutEngine()
        
        # 1. Parse resume
        result = parser.parse_file(sample_docx)
        bullets = result['raw_content'][:3]  # First 3 lines
        
        # 2. Analyze JD
        jd_text = "Looking for Python developer with FastAPI experience"
        jd_result = jd_agent.analyze_jd(jd_text)
        keywords = jd_result['keywords']
        
        # 3. Rewrite bullets
        constraints = [{"max_chars": 100} for _ in bullets]
        rewritten = rewrite_agent.batch_rewrite(bullets, keywords, constraints, "safe")
        
        # 4. Validate
        for i, new_bullet in enumerate(rewritten):
            is_valid = engine.validate_constraints(bullets[i], new_bullet, 100)
            assert is_valid is True or len(new_bullet) == len(bullets[i])
