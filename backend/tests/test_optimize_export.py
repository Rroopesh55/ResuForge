import io
import json
import tempfile
import uuid
from pathlib import Path

import docx
import pytest
from fastapi.testclient import TestClient

from backend.main import app


client = TestClient(app)


generated_pdf_paths: list[Path] = []


@pytest.fixture(autouse=True)
def stub_external_agents(monkeypatch, tmp_path):
    generated_pdf_paths.clear()

    def fake_batch_rewrite(bullets, keywords, constraints, style="safe"):
        return [f"{text} [{style}]" for text in bullets]

    def fake_convert_docx_to_pdf(docx_path: str):
        pdf_path = tmp_path / f"{Path(docx_path).stem}_{uuid.uuid4().hex}.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 stub")
        generated_pdf_paths.append(pdf_path)
        return str(pdf_path)

    monkeypatch.setattr("backend.main.rewrite_agent.batch_rewrite", fake_batch_rewrite)
    monkeypatch.setattr("backend.main.pdf_converter_agent.convert_docx_to_pdf", fake_convert_docx_to_pdf)


def _create_sample_docx(tmp_dir: Path) -> Path:
    document = docx.Document()
    document.add_paragraph("Designed APIs for fintech workflows")
    document.add_paragraph("Improved latency by 40% across services")
    path = tmp_dir / "resume.docx"
    document.save(path)
    return path


def test_optimize_returns_validated_bullets():
    payload = {
        "bullets": [
            "Built data pipelines for analytics",
            "Reduced costs by 30 percent",
        ],
        "keywords": ["analytics", "cost"],
        "constraints": [{"max_chars": 200}, {"max_chars": 200}],
        "style": "bold",
    }

    response = client.post("/optimize", json=payload)
    assert response.status_code == 200

    data = response.json()
    assert "rewritten_bullets" in data
    assert "validation" in data
    assert len(data["rewritten_bullets"]) == len(payload["bullets"])
    assert len(data["validation"]) == len(payload["bullets"])


def test_optimize_and_export_docx_roundtrip(tmp_path: Path):
    original_path = _create_sample_docx(tmp_path)
    paragraphs = [
        "Designed APIs for fintech workflows",
        "Improved latency by 40% across services",
    ]
    constraints = [{"max_chars": len(p) + 50} for p in paragraphs]

    with original_path.open("rb") as fh:
        files = {
            "file": ("resume.docx", fh, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }
        data = {
            "bullets_json": json.dumps(paragraphs),
            "constraints_json": json.dumps(constraints),
            "keywords_json": json.dumps(["fintech", "latency"]),
            "style": "safe",
            "output_format": "docx",
        }
        response = client.post("/optimize-and-export", files=files, data=data)

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    validation_header = response.headers.get("x-validation-results")
    assert validation_header is not None
    assert json.loads(validation_header) == [True, True]

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
        temp_docx.write(response.content)
        temp_docx_path = Path(temp_docx.name)

    try:
        rebuilt = docx.Document(temp_docx_path)
        exported_text = [para.text for para in rebuilt.paragraphs if para.text.strip()]
        assert exported_text, "Exported DOCX should contain text"
        assert len(exported_text) == len(paragraphs)
        for original, rewritten in zip(paragraphs, exported_text):
            assert rewritten.startswith(original)
            assert rewritten.endswith("[safe]")
    finally:
        temp_docx_path.unlink(missing_ok=True)


def test_optimize_and_export_respects_final_bullets(tmp_path: Path):
    original_path = _create_sample_docx(tmp_path)
    original_paragraphs = [
        "Designed APIs for fintech workflows",
        "Improved latency by 40% across services",
    ]
    final_paragraphs = ["Custom 1", "Custom 2 with keywords"]
    constraints = [{"max_chars": len(p) + 50} for p in original_paragraphs]

    with original_path.open("rb") as fh:
        files = {
            "file": ("resume.docx", fh, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }
        data = {
            "bullets_json": json.dumps(original_paragraphs),
            "constraints_json": json.dumps(constraints),
            "keywords_json": json.dumps(["fintech", "latency"]),
            "style": "safe",
            "output_format": "docx",
            "final_bullets_json": json.dumps(final_paragraphs),
        }
        response = client.post("/optimize-and-export", files=files, data=data)

    assert response.status_code == 200
    validation_header = response.headers.get("x-validation-results")
    assert validation_header is not None
    assert json.loads(validation_header) == [True, True]

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
        temp_docx.write(response.content)
        temp_docx_path = Path(temp_docx.name)

    try:
        rebuilt = docx.Document(temp_docx_path)
        exported_text = [para.text for para in rebuilt.paragraphs if para.text.strip()]
        assert exported_text == final_paragraphs
    finally:
        temp_docx_path.unlink(missing_ok=True)


def test_optimize_and_export_pdf_cleans_temp_files(tmp_path: Path):
    original_path = _create_sample_docx(tmp_path)
    paragraphs = [
        "Designed APIs for fintech workflows",
        "Improved latency by 40% across services",
    ]
    constraints = [{"max_chars": len(p) + 50} for p in paragraphs]

    with original_path.open("rb") as fh:
        files = {
            "file": ("resume.docx", fh, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }
        data = {
            "bullets_json": json.dumps(paragraphs),
            "constraints_json": json.dumps(constraints),
            "keywords_json": json.dumps(["fintech", "latency"]),
            "style": "safe",
            "output_format": "pdf",
        }
        response = client.post("/optimize-and-export", files=files, data=data)

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/pdf"
    validation_header = response.headers.get("x-validation-results")
    assert validation_header is not None
    assert json.loads(validation_header) == [True, True]
    assert generated_pdf_paths, "PDF path should have been recorded"
    assert not generated_pdf_paths[-1].exists()


def test_convert_to_pdf_streams_pdf_and_removes_temp(tmp_path: Path):
    resume_path = _create_sample_docx(tmp_path)

    with resume_path.open("rb") as fh:
        files = {"file": ("resume.docx", fh, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        response = client.post("/convert-to-pdf", files=files)

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/pdf"
    assert generated_pdf_paths, "PDF path should have been recorded"
    assert not generated_pdf_paths[-1].exists()


def test_optimize_and_export_reports_invalid_overrides(tmp_path: Path):
    original_path = _create_sample_docx(tmp_path)
    paragraphs = [
        "Designed APIs for fintech workflows",
        "Improved latency by 40% across services",
    ]
    constraints = [
        {"max_chars": len(paragraphs[0])},
        {"max_chars": len(paragraphs[1]) + 100},
    ]
    final_override = [
        paragraphs[0] + " plus extra details that exceed constraints",
        paragraphs[1],
    ]

    with original_path.open("rb") as fh:
        files = {
            "file": ("resume.docx", fh, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        }
        data = {
            "bullets_json": json.dumps(paragraphs),
            "constraints_json": json.dumps(constraints),
            "keywords_json": json.dumps(["fintech", "latency"]),
            "style": "safe",
            "output_format": "docx",
            "final_bullets_json": json.dumps(final_override),
        }
        response = client.post("/optimize-and-export", files=files, data=data)

    assert response.status_code == 200
    validation_header = response.headers.get("x-validation-results")
    assert validation_header is not None
    assert json.loads(validation_header) == [False, True]

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
        temp_docx.write(response.content)
        temp_docx_path = Path(temp_docx.name)

    try:
        rebuilt = docx.Document(temp_docx_path)
        exported_text = [para.text for para in rebuilt.paragraphs if para.text.strip()]
        assert exported_text[0] == paragraphs[0]
        assert exported_text[1] == paragraphs[1]
    finally:
        temp_docx_path.unlink(missing_ok=True)

