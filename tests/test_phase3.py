import pytest
from fastapi.testclient import TestClient
from backend.main import app
import os
from docx import Document

client = TestClient(app)

def test_validate_edit_valid():
    response = client.post("/validate-edit", json={
        "original_text": "Hello world",
        "new_text": "Hello there",
        "max_chars": 20
    })
    assert response.status_code == 200
    data = response.json()
    assert data["valid"] is True
    assert data["diff"] == 0

def test_validate_edit_invalid():
    response = client.post("/validate-edit", json={
        "original_text": "Hello world",
        "new_text": "This is way too long for the constraint",
        "max_chars": 10
    })
    assert response.status_code == 200
    data = response.json()
    assert data["valid"] is False
    assert data["diff"] > 0

def test_update_content_flow():
    # 1. Create a dummy DOCX
    doc = Document()
    doc.add_paragraph("Original text content")
    test_docx = "test_phase3.docx"
    doc.save(test_docx)
    
    try:
        # 2. Call update-content
        with open(test_docx, "rb") as f:
            response = client.post(
                "/update-content",
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data={
                    "original_text": "Original text content",
                    "new_text": "Updated text content"
                }
            )
            
        assert response.status_code == 200
        assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        # 3. Save the result and verify it's a valid DOCX
        output_docx = "test_phase3_output.docx"
        with open(output_docx, "wb") as f:
            f.write(response.content)
            
        # Verify content changed
        new_doc = Document(output_docx)
        full_text = "\n".join([p.text for p in new_doc.paragraphs])
        assert "Updated text content" in full_text
        assert "Original text content" not in full_text
        
    finally:
        # Cleanup
        if os.path.exists(test_docx):
            os.remove(test_docx)
        if os.path.exists(output_docx):
            os.remove(output_docx)
